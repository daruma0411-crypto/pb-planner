"""
PB企画支援チャットボット — Flask Webサーバー
=============================================
GET  /            → static/index.html
POST /api/chat    → PB企画チャット（Claude Sonnet + Function Calling）
GET  /api/health  → ヘルスチェック
GET  /api/download/<filename> → 生成ファイルダウンロード
"""
import glob as glob_module
import json
import os
import re
import time
import sys
import tempfile
import traceback

from dotenv import load_dotenv
from flask import (Flask, Response, abort, jsonify, request,
                   send_from_directory, send_file)

load_dotenv()

CLAUDE_API_KEY = (os.environ.get('ANTHROPIC_API_KEY')
                  or os.environ.get('CLAUDE_API_KEY')
                  or os.environ.get('API_KEY')
                  or '')

app = Flask(__name__, static_folder='static')

# workspace/data/ のルートパス
_WORKSPACE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'workspace')

# ダウンロード用一時ディレクトリ
_DOWNLOADS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'downloads')
os.makedirs(_DOWNLOADS_DIR, exist_ok=True)

# テンプレートディレクトリ
_TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')

# ================================================================
# セッションストレージ（ファイル永続化 / Redis / フォールバック: メモリ内）
# ================================================================
_SESSION_TTL = 3600  # 1時間

_redis_client = None
_REDIS_URL = os.environ.get('REDIS_URL', '')
if _REDIS_URL:
    try:
        import redis
        _redis_client = redis.from_url(_REDIS_URL, decode_responses=True)
        _redis_client.ping()
        print("[SESSION] Redis connected", flush=True, file=sys.stderr)
    except Exception as _e:
        print(f"[SESSION] Redis connect failed: {_e}, falling back to file",
              flush=True, file=sys.stderr)
        _redis_client = None
else:
    print("[SESSION] No REDIS_URL, using file-based sessions",
          flush=True, file=sys.stderr)

# ファイルベースセッション用ディレクトリ
_SESSION_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.sessions')
os.makedirs(_SESSION_DIR, exist_ok=True)

_SESSIONS = {}  # メモリキャッシュ（高速アクセス用）


# ================================================================
# セッション管理
# ================================================================

def _new_session_dict():
    """空のセッション辞書を返す"""
    return {
        'history': [],
        'pb_card': {
            'asone_part_no': None,
            'price': None,
            'jan_code': None,
            'maker_part_no': None,
            'quantity': None,
            'catchcopy': None,
            'spec_diff': None,
        },
        'base_product': None,
        'framework_results': {},
        'last_search_results': [],
        'confirmed_specs': [],  # 確定済み仕様諸元リスト [{no, name, value}, ...]
    }


def _session_filepath(session_id):
    """セッションIDからファイルパスを生成（安全なファイル名に変換）"""
    safe_id = re.sub(r'[^a-zA-Z0-9_-]', '_', session_id)[:100]
    return os.path.join(_SESSION_DIR, f'{safe_id}.json')


def get_or_create_session(session_id):
    """セッション取得 or 新規作成"""
    # 1. Redis
    if _redis_client:
        try:
            raw = _redis_client.get(f'pb:{session_id}')
            if raw:
                return json.loads(raw)
        except Exception:
            pass

    # 2. メモリキャッシュ
    if session_id in _SESSIONS:
        return _SESSIONS[session_id]

    # 3. ファイル（再デプロイ後の復元）
    fpath = _session_filepath(session_id)
    if os.path.exists(fpath):
        try:
            with open(fpath, 'r', encoding='utf-8') as f:
                session = json.loads(f.read())
            _SESSIONS[session_id] = session  # メモリキャッシュに戻す
            print(f"[SESSION] Restored from file: {session_id}", flush=True, file=sys.stderr)
            return session
        except Exception as e:
            print(f"[SESSION] File read error: {e}", flush=True, file=sys.stderr)

    return _new_session_dict()


def save_session(session_id, session):
    """セッション保存（Redis + ファイル + メモリ）"""
    # Redis
    if _redis_client:
        try:
            _redis_client.setex(f'pb:{session_id}', _SESSION_TTL,
                                json.dumps(session, ensure_ascii=False))
        except Exception:
            pass

    # メモリキャッシュ
    _SESSIONS[session_id] = session

    # ファイル永続化（再デプロイ対策）
    try:
        fpath = _session_filepath(session_id)
        with open(fpath, 'w', encoding='utf-8') as f:
            f.write(json.dumps(session, ensure_ascii=False))
    except Exception as e:
        print(f"[SESSION] File save error: {e}", flush=True, file=sys.stderr)

    # 古いセッションファイルの掃除（100件超で古い順に削除）
    try:
        files = sorted(
            [os.path.join(_SESSION_DIR, f) for f in os.listdir(_SESSION_DIR) if f.endswith('.json')],
            key=os.path.getmtime
        )
        while len(files) > 100:
            os.remove(files.pop(0))
    except Exception:
        pass


# ================================================================
# 製品データロード
# ================================================================

def _load_all_products():
    """workspace/data/*/products.jsonl から全製品を読み込む"""
    products = []
    jsonl_paths = glob_module.glob(
        os.path.join(_WORKSPACE_DIR, 'data', '*', 'products.jsonl')
    )
    for path in jsonl_paths:
        try:
            with open(path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line:
                        products.append(json.loads(line))
        except Exception:
            pass
    return products


def _search_products(query=None, category=None, maker=None, usage=None):
    """製品DB検索"""
    all_products = _load_all_products()
    results = []
    for p in all_products:
        # カテゴリフィルタ
        if category and p.get('category', '').lower() != category.lower():
            continue
        # メーカーフィルタ
        if maker and maker.lower() not in p.get('maker', '').lower():
            continue
        # 用途フィルタ
        if usage and usage.lower() not in (p.get('usage') or '').lower():
            continue
        # テキスト検索
        if query:
            q_lower = query.lower()
            searchable = ' '.join([
                p.get('name', ''),
                p.get('model', ''),
                p.get('maker', ''),
                p.get('category', ''),
                p.get('description', ''),
                json.dumps(p.get('specs', {}), ensure_ascii=False),
            ]).lower()
            if q_lower not in searchable:
                continue
        results.append(p)
    return results


# ================================================================
# Function Calling ツール定義
# ================================================================

FC_TOOLS = [
    {
        "name": "search_products",
        "description": "仕入れ先製品データベースを検索する。カテゴリ、メーカー、テキスト検索、用途で絞り込み可能。",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "検索キーワード（製品名、型番、スペックなど）"
                },
                "category": {
                    "type": "string",
                    "description": "製品カテゴリ（autoclave, accessory, consumable等）"
                },
                "maker": {
                    "type": "string",
                    "description": "メーカー名"
                },
                "usage": {
                    "type": "string",
                    "description": "用途（ラボ, メディカル, 調理等）"
                }
            }
        }
    },
    {
        "name": "set_pb_field",
        "description": "PB企画カードの項目を確定する。field_nameとvalueを指定。",
        "input_schema": {
            "type": "object",
            "properties": {
                "field_name": {
                    "type": "string",
                    "enum": ["asone_part_no", "price", "jan_code",
                             "maker_part_no", "quantity", "catchcopy", "spec_diff"],
                    "description": "確定する項目名"
                },
                "value": {
                    "type": "string",
                    "description": "確定する値"
                }
            },
            "required": ["field_name", "value"]
        }
    },
    {
        "name": "get_pb_card",
        "description": "現在のPB企画カードの状態を取得する。",
        "input_schema": {
            "type": "object",
            "properties": {}
        }
    },
    {
        "name": "analyze_framework",
        "description": "フレームワーク分析を実行する。3C分析、SWOT分析、ポジショニングマップ、5Forces分析、価格帯マップから選択。ポジショニングマップでは軸を指定可能（未指定時は軸候補を返す）。",
        "input_schema": {
            "type": "object",
            "properties": {
                "framework": {
                    "type": "string",
                    "enum": ["3c", "swot", "positioning", "5forces", "price_map"],
                    "description": "実行するフレームワーク"
                },
                "axis_x": {
                    "type": "string",
                    "description": "ポジショニングX軸（positioning時のみ）。例: 価格, 缶体容量, 機能数"
                },
                "axis_y": {
                    "type": "string",
                    "description": "ポジショニングY軸（positioning時のみ）。例: 缶体容量, コース数, 本体質量"
                }
            },
            "required": ["framework"]
        }
    },
    {
        "name": "generate_pim_excel",
        "description": "PIMデータExcelファイルを生成する。PBカードの全項目が確定済みである必要がある。",
        "input_schema": {
            "type": "object",
            "properties": {}
        }
    },
    {
        "name": "generate_proposal_word",
        "description": "企画書Wordファイルを生成する。PBカードとフレームワーク分析結果を含む。",
        "input_schema": {
            "type": "object",
            "properties": {}
        }
    },
    {
        "name": "translate_to_english",
        "description": "PIMデータを英訳してExcelファイルを生成する。",
        "input_schema": {
            "type": "object",
            "properties": {}
        }
    },
    {
        "name": "generate_catalog_html",
        "description": "カタログHTMLファイルを生成する。ベース製品情報とPBカード情報を含む。",
        "input_schema": {
            "type": "object",
            "properties": {}
        }
    },
    {
        "name": "save_spec_list",
        "description": "仕様諸元リストを保存する。連番つき仕様一覧をユーザーに提示した後、そのリストをセッションに保存する。企画書Word最終ページの仕様諸元表に使われる。",
        "input_schema": {
            "type": "object",
            "properties": {
                "specs": {
                    "type": "array",
                    "description": "仕様リスト。各要素は {no: 番号, name: 項目名, value: 値} の形式",
                    "items": {
                        "type": "object",
                        "properties": {
                            "no": {"type": "integer", "description": "連番"},
                            "name": {"type": "string", "description": "項目名"},
                            "value": {"type": "string", "description": "値"}
                        },
                        "required": ["no", "name", "value"]
                    }
                }
            },
            "required": ["specs"]
        }
    },
    {
        "name": "update_spec_item",
        "description": "仕様諸元リストの特定の番号の項目を更新する。「8番を300に変更」のような指示に対応。変更はspec_changesに自動記録される。",
        "input_schema": {
            "type": "object",
            "properties": {
                "no": {
                    "type": "integer",
                    "description": "更新する項目の連番"
                },
                "value": {
                    "type": "string",
                    "description": "新しい値。項目名を変更する場合はname引数も指定。"
                },
                "name": {
                    "type": "string",
                    "description": "新しい項目名（省略時は項目名は変更しない）"
                },
                "reason": {
                    "type": "string",
                    "description": "変更理由（ユーザーの指示内容を要約）"
                }
            },
            "required": ["no", "value"]
        }
    },
    {
        "name": "delete_spec_item",
        "description": "仕様諸元リストから特定の番号の項目を削除する。「15番は不要」のような指示に対応。変更はspec_changesに自動記録される。",
        "input_schema": {
            "type": "object",
            "properties": {
                "no": {
                    "type": "integer",
                    "description": "削除する項目の連番"
                },
                "reason": {
                    "type": "string",
                    "description": "削除理由（ユーザーの指示内容を要約）"
                }
            },
            "required": ["no"]
        }
    },
    {
        "name": "add_spec_item",
        "description": "仕様諸元リストに新しい項目を追加する。末尾に連番を振って追加。変更はspec_changesに自動記録される。",
        "input_schema": {
            "type": "object",
            "properties": {
                "name": {
                    "type": "string",
                    "description": "追加する項目名"
                },
                "value": {
                    "type": "string",
                    "description": "値"
                },
                "reason": {
                    "type": "string",
                    "description": "追加理由（ユーザーの指示内容を要約）"
                }
            },
            "required": ["name", "value"]
        }
    },
]


# ================================================================
# システムプロンプト
# ================================================================

_PB_CONSULTANT_SYSTEM_PROMPT = """\
あなたはアズワンPB企画の専門コンサルタントです。仕入れ先の製品データを元に、PB化の壁打ちパートナーとしてユーザーと進めます。

## 最重要ルール：会話の文脈を読め

あなたの最大の仕事は「ユーザーの話を聞くこと」。以下を絶対に守れ：

1. **聞き返すな、動け。** ユーザーの意図が文脈から分かるなら即座にツールを呼べ。「確定しますか？」「よろしいですか？」は禁止。
2. **前の会話を忘れるな。** 会話で出た製品名・型番・価格・分析結果はすべて覚えていること。「先に製品を確定してください」と言い返すのは最悪の対応。
3. **一度に処理しろ。** ユーザーが「品番：200003003、JAN：9000021、入数：1」と言ったら、3つとも一度にset_pb_fieldで確定。1つだけ確定して残りを聞き返すのは禁止。
4. **既に確定した値を忘れるな。** PBカードに既に入っている値は触るな。新しく指示された値だけ追加・更新しろ。
5. **ユーザーの指示した書式に従え。** 「連番ふって」→ 1,2,3...の通し番号を全項目に振れ（カテゴリで途切れさせるな）。「表にして」→ テーブル形式。「一覧」→ 全項目を漏れなく。ユーザーが後から「8番を変更」と言えるように番号は必ず振ること。

## 具体的なNG/OK例

NG: 「3C分析を実行するには、まずベース製品を確定する必要があります」
OK: （会話で出ている型番でsearch_products→analyze_framework即実行）

NG: 「FLS-1000をベースにしますか？確定していただければ分析を実行します」
OK: （FLS-1000の話をしているのだからsearch_productsしてそのまま分析実行）

NG: 「価格を90万円で確定しますか？」
OK: （「90万円で」と言われたらset_pb_field即実行）

NG: 品番設定後に「残り4項目を決めましょう！」と言って既に設定済みの項目を未確定表示
OK: 既に確定済みの項目はそのまま維持し、新しく設定した項目だけ報告

NG: 「連番ふって」と言われてカテゴリ見出し+箇条書きで返す（番号なし）
OK: 1. ○○ / 2. △△ / 3. □□ ... と全項目に通し番号を振る。ユーザーが「8番を変更して」と言えるようにする

NG: ユーザーが「①と②に加えて別項目で後２つ」と言ったのに全部で6つ出した後、番号をリセットして1から振り直す
OK: 既に出した①②③④に続けて⑤⑥として追加。番号体系を維持する

## データルール
- search_productsの結果のみ使用（ハルシネーション厳禁）
- DBに無い情報は「データベースに登録がありません」
- PBカード値の確定は必ずset_pb_fieldツールを使用

## PB企画カードの7項目
- asone_part_no: アズワン品番
- price: 販売価格（税抜）
- jan_code: JANコード
- maker_part_no: メーカー型番
- quantity: 入数
- catchcopy: キャッチコピー
- spec_diff: 仕様差分（※手動設定不要。spec_changesから企画書生成時に自動生成される）

## ベース製品の扱い
- ユーザーが型番名を出した時点でsearch_productsで検索し、ベース製品として扱え
- 「FLS-1000の特徴を教えて」→ search_productsしてFLS-1000の情報で回答。これがベース製品。
- 以後「比較して」「3C分析して」→ そのベース製品で即実行。「確定しますか？」不要。
- maker_part_noは最初の検索時にset_pb_fieldで自動設定してよい

## 製品紹介のルール
製品の特徴を聞かれたら、以下の順で紹介：
1. **設計思想** (design_concept): 誰のために、なぜ作られたか
2. **製品特長** (features): 使いやすさ、独自技術、特許機能
3. **運転コース** (operation_courses): 使い方のバリエーション
4. **安全機能** (safety): 概要のみ
5. **オプション** (options): 拡張性
6. **収納量** (storage_capacity): 実用データ
7. **スペック表** (specs): 数値は最後

スペック表の羅列だけは不可。製品の「ストーリー」として語れ。

## 仕様諸元リスト（ストック型ワークフロー）

### 基本フロー
1. ユーザーが「仕様を連番で並べて」→ 全仕様に通し番号をつけて表示 → 直後に **save_spec_list** で保存（確認不要）
2. 壁打ち中にユーザーが指示を出す → 変更を即適用＆ログ記録：
   - 「8番をΦ460に変更」→ **update_spec_item**（変更ログ自動記録）
   - 「15番は不要」→ **delete_spec_item**（削除ログ自動記録）
   - 「消費電力の項目を追加して」→ **add_spec_item**（追加ログ自動記録）
3. すべての壁打ちが終わったら企画書を生成 → 統合された仕様諸元表 + 変更履歴が含まれる

### 重要ルール
- 変更指示は必ずツールで実行（口頭メモではなくデータとして記録）
- 「仕様差分(spec_diff)」はPBカードに手動設定しなくてよい（spec_changesから自動生成）
- 企画書Wordの構成: フレームワーク分析 → PB製品仕様 → 仕様諸元表（最終ページ）
- 諸元表 = ベースfix仕様 + 壁打ちで蓄積された全変更を統合した最終版

## 壁打ちの進め方
- ユーザーが迷ったら、データで選択肢を絞る
- 価格設定は原価率・競合観点でアドバイス
- キャッチコピーは差別化ポイントから案を提示
- フレームワーク分析はいつでも実行可能（会話中に製品が特定されていれば）
- **ポジショニング分析のフロー（必ず守ること）**:
  1. まずanalyze_framework(framework="positioning")をaxis_x/axis_y未指定で呼ぶ
  2. needs_axes=Trueが返るので、その中のaxis_candidatesとsuggested_pairsをユーザーに提示する
  3. 自分で軸を決めてはいけない。必ずユーザーに選んでもらう
  4. ユーザーが軸を選んだら、**必ず** analyze_framework(framework="positioning", axis_x="...", axis_y="...")で再度ツールを呼ぶ。自分でテキスト回答を生成してはいけない
  5. axis_xとaxis_yにはsuggested_pairsで提案した軸名をそのまま使う
- **重要**: フレームワーク分析（3C/SWOT/5Forces/価格帯マップ/ポジショニング）を実行する際は、必ずanalyze_frameworkツールを呼び出すこと。ツールを呼ばずに自分で分析テキストを生成してはいけない。ツールを呼ぶことでチャートやグラフが自動表示される
- フレームワーク分析は競合DBの実データに基づく。推測・ハルシネーションではなくファクトベースで分析する
"""


def _build_system_prompt(session):
    """セッション状態に応じてシステムプロンプトを動的構築"""
    parts = [_PB_CONSULTANT_SYSTEM_PROMPT]

    # PBカード状態
    pb = session.get('pb_card', {})
    filled = sum(1 for v in pb.values() if v is not None)
    total = len(pb)
    parts.append(f"\n## 現在のPBカード状態 ({filled}/{total}項目確定)")
    for k, v in pb.items():
        status = f"✅ {v}" if v else "⏳ 未確定"
        parts.append(f"- {k}: {status}")

    # ベース製品情報
    base = session.get('base_product')
    if base:
        parts.append(f"\n## ベース製品")
        parts.append(f"- 製品名: {base.get('name', '不明')}")
        parts.append(f"- メーカー: {base.get('maker', '不明')}")
        parts.append(f"- 型番: {base.get('model', '不明')}")
        if base.get('price'):
            parts.append(f"- 価格: {base.get('price')}")
        if base.get('description'):
            parts.append(f"- 概要: {base.get('description')}")
        if base.get('design_concept'):
            parts.append(f"- 設計思想: {base.get('design_concept')}")
        if base.get('features'):
            parts.append(f"- 製品特長: {json.dumps(base['features'], ensure_ascii=False)}")
        if base.get('operation_courses'):
            parts.append(f"- 運転コース: {json.dumps(base['operation_courses'], ensure_ascii=False)}")
        if base.get('options'):
            parts.append(f"- オプション: {json.dumps(base['options'], ensure_ascii=False)}")
        if base.get('storage_capacity'):
            parts.append(f"- 収納量: {json.dumps(base['storage_capacity'], ensure_ascii=False)}")
        if base.get('specs'):
            parts.append(f"- スペック: {json.dumps(base['specs'], ensure_ascii=False)}")

    # 分析済みフレームワーク
    fw = session.get('framework_results', {})
    if fw:
        parts.append(f"\n## 分析済みフレームワーク: {', '.join(fw.keys())}")

    # 仕様諸元リスト・変更ログ状態
    specs = session.get('confirmed_specs', [])
    changes = session.get('spec_changes', [])
    if specs:
        parts.append(f"\n## 仕様諸元リスト: {len(specs)}項目保存済み")
    if changes:
        parts.append(f"## 仕様変更ログ: {len(changes)}件蓄積")
        for ch in changes[-5:]:  # 直近5件
            t = ch.get('type')
            if t == 'update':
                parts.append(f"  - 変更: {ch.get('name')} → {ch.get('new')}")
            elif t == 'delete':
                parts.append(f"  - 削除: {ch.get('name')}")
            elif t == 'add':
                parts.append(f"  - 追加: {ch.get('name')}")

    return '\n'.join(parts)


# ================================================================
# ツールハンドラ
# ================================================================

def handle_search_products(args, session):
    """製品DB検索（1件ヒット時は自動的にベース製品としてリンク）"""
    results = _search_products(
        query=args.get('query'),
        category=args.get('category'),
        maker=args.get('maker'),
        usage=args.get('usage'),
    )
    session['last_search_results'] = results[:20]

    if not results:
        return {"found": 0, "message": "該当する製品が見つかりませんでした。"}

    # 型番完全一致 or 1件のみヒット → ベース製品として自動リンク
    query = args.get('query', '')
    auto_linked = None
    if query:
        for p in results:
            if p.get('model', '').upper() == query.upper():
                session['base_product'] = p
                auto_linked = p.get('model')
                break
    if not auto_linked and len(results) == 1:
        session['base_product'] = results[0]
        auto_linked = results[0].get('model')

    # 結果フォーマット
    items = []
    for p in results[:20]:
        item = {
            "id": p.get('id', ''),
            "name": p.get('name', ''),
            "model": p.get('model', ''),
            "maker": p.get('maker', ''),
            "category": p.get('category', ''),
            "price": p.get('price', ''),
            "usage": p.get('usage', ''),
        }
        if p.get('specs'):
            item['specs'] = p['specs']
        if p.get('description'):
            item['description'] = p['description']
        # リッチ製品データ（特長・オプション等）
        for rich_key in ('design_concept', 'features', 'operation_courses',
                         'safety', 'options', 'storage_capacity', 'maintenance',
                         'regulatory_note'):
            if p.get(rich_key):
                item[rich_key] = p[rich_key]
        items.append(item)

    result = {"found": len(results), "showing": len(items), "products": items}
    if auto_linked:
        result["auto_linked_base_product"] = auto_linked
    return result


def handle_set_pb_field(args, session):
    """PBカード項目確定"""
    field = args.get('field_name')
    value = args.get('value')

    valid_fields = ['asone_part_no', 'price', 'jan_code',
                    'maker_part_no', 'quantity', 'catchcopy', 'spec_diff']
    if field not in valid_fields:
        return {"error": f"無効な項目名: {field}", "valid_fields": valid_fields}

    session['pb_card'][field] = value

    # maker_part_noが設定されたらベース製品を自動リンク
    if field == 'maker_part_no' and value:
        found = False
        # まずlast_search_resultsから探す
        for p in session.get('last_search_results', []):
            if p.get('model') == value:
                session['base_product'] = p
                found = True
                break
        # 見つからなければDBを直接検索
        if not found:
            db_results = _search_products(query=value)
            for p in db_results:
                if p.get('model') == value:
                    session['base_product'] = p
                    found = True
                    break
            if not found and db_results:
                session['base_product'] = db_results[0]

    filled = sum(1 for v in session['pb_card'].values() if v is not None)
    total = len(session['pb_card'])
    return {
        "success": True,
        "field": field,
        "value": value,
        "progress": f"{filled}/{total}",
        "all_filled": filled == total,
    }


def handle_get_pb_card(args, session):
    """PBカード状態取得"""
    pb = session.get('pb_card', {})
    filled = sum(1 for v in pb.values() if v is not None)
    total = len(pb)
    return {
        "pb_card": pb,
        "progress": f"{filled}/{total}",
        "all_filled": filled == total,
        "base_product": session.get('base_product'),
    }


def _build_competitor_summary(category_products, base_product=None):
    """競合製品のファクトサマリを構築"""
    by_maker = {}
    for p in category_products:
        maker = p.get('maker', '不明')
        if maker not in by_maker:
            by_maker[maker] = []
        by_maker[maker].append(p)

    lines = []
    for maker, products in sorted(by_maker.items()):
        prices = [p.get('price_numeric') for p in products if p.get('price_numeric')]
        models = [p.get('model', '') for p in products]
        price_range = f"¥{min(prices):,}〜¥{max(prices):,}" if prices else "価格不明"
        lines.append(f"- **{maker}** ({len(products)}機種): {', '.join(models[:5])}{'...' if len(models) > 5 else ''} / {price_range}")
    return '\n'.join(lines)


def _build_positioning_data(category_products, axis_x, axis_y):
    """軸に基づくポジショニングデータを構築"""
    data_points = []
    for p in category_products:
        specs = p.get('specs', {})
        x_val = _extract_numeric_spec(p, axis_x)
        y_val = _extract_numeric_spec(p, axis_y)
        data_points.append({
            'maker': p.get('maker', ''),
            'model': p.get('model', ''),
            'x_value': x_val,
            'y_value': y_val,
            'x_label': _extract_spec_text(p, axis_x),
            'y_label': _extract_spec_text(p, axis_y),
        })
    return data_points


def _extract_numeric_spec(product, axis_name):
    """製品から軸に対応する数値を抽出"""
    import re as _re
    if '価格' in axis_name or 'price' in axis_name.lower():
        return product.get('price_numeric')
    specs = product.get('specs', {})
    # 類義語マッピング（キー名のゆらぎ吸収）
    _synonyms = {
        '重量': ['重さ', '質量', 'weight', '重 量'],
        '容量': ['有効容量', '内容量', '缶体容量'],
        '温度': ['使用温度', '温度範囲', '使用温度範囲'],
    }
    search_terms = [axis_name]
    for base_word, syns in _synonyms.items():
        if base_word in axis_name:
            search_terms.extend(syns)
        for s in syns:
            if s in axis_name:
                search_terms.append(base_word)
                search_terms.extend(syns)
                break
    # スペックキーから検索（部分一致、類義語対応）
    for key, val in specs.items():
        for term in search_terms:
            if term in key or key in term:
                nums = _re.findall(r'[\d.]+', str(val).replace(',', ''))
                if nums:
                    try:
                        return float(nums[0])
                    except ValueError:
                        pass
    return None


def _extract_spec_text(product, axis_name):
    """製品から軸に対応するテキスト値を抽出"""
    if '価格' in axis_name or 'price' in axis_name.lower():
        return product.get('price', '')
    specs = product.get('specs', {})
    for key, val in specs.items():
        if axis_name in key:
            return str(val)
    return ''


def _get_axis_candidates(category_products):
    """競合DB内のスペックキーから軸候補を生成"""
    spec_keys = {}
    for p in category_products:
        for key, val in p.get('specs', {}).items():
            if key not in spec_keys:
                spec_keys[key] = 0
            spec_keys[key] += 1
    # 出現頻度順でソート、数値を含むものを優先
    import re as _re
    candidates = []
    for key, count in sorted(spec_keys.items(), key=lambda x: -x[1]):
        if count >= 3:
            candidates.append(key)
    # 価格は常に候補に含める
    price_count = sum(1 for p in category_products if p.get('price_numeric'))
    if price_count >= 2:
        candidates.insert(0, '価格')
    return candidates[:10]


# ================================================================
# フレームワーク分析ビジュアルデータ生成
# ================================================================

_MAKER_COLORS = {
    'トミー精工': {'bg': 'rgba(33, 150, 243, 0.7)', 'border': '#2196F3'},
    'ヒラヤマ': {'bg': 'rgba(244, 67, 54, 0.7)', 'border': '#F44336'},
    'ヤマト科学': {'bg': 'rgba(76, 175, 80, 0.7)', 'border': '#4CAF50'},
    'アルプ': {'bg': 'rgba(156, 39, 176, 0.7)', 'border': '#9C27B0'},
}
_DEFAULT_COLOR = {'bg': 'rgba(158, 158, 158, 0.7)', 'border': '#9E9E9E'}


def _generate_framework_visual(framework, category_products, base_product, **kwargs):
    """フレームワーク分析のビジュアルデータを生成"""
    print(f"[VISUAL] Generating visual for: {framework}, products: {len(category_products)}, base: {base_product.get('model') if base_product else 'None'}", flush=True)
    try:
        result = None
        if framework == 'swot':
            result = _gen_swot_visual(category_products, base_product)
        elif framework == 'positioning':
            result = _gen_positioning_visual(
                category_products, base_product,
                kwargs.get('axis_x', ''), kwargs.get('axis_y', ''))
        elif framework == '5forces':
            result = _gen_5forces_visual(category_products, base_product)
        elif framework == 'price_map':
            result = _gen_price_map_visual(category_products, base_product)
        print(f"[VISUAL] Result for {framework}: {'OK (type=' + result.get('type', '?') + ')' if result else 'None'}", flush=True)
        return result
    except Exception as e:
        import traceback
        print(f"[VISUAL] Error generating {framework}: {e}\n{traceback.format_exc()}", flush=True)
    return None


def _gen_swot_visual(category_products, base_product):
    """SWOT 2x2グリッド用データ"""
    base_maker = base_product.get('maker', '')
    base_specs = base_product.get('specs', {})
    base_price = base_product.get('price_numeric')

    comp_prices = [p.get('price_numeric') for p in category_products
                   if p.get('price_numeric') and p.get('maker') != base_maker]
    all_prices = [p.get('price_numeric') for p in category_products if p.get('price_numeric')]
    makers = set(p.get('maker', '') for p in category_products)

    # Strengths
    strengths = []
    cap_text = base_specs.get('有効容量', base_specs.get('缶体有効内容積', ''))
    if cap_text:
        strengths.append(f'有効容量 {cap_text}')
    temp_text = base_specs.get('使用温度範囲', base_specs.get('滅菌温度範囲', ''))
    if temp_text:
        strengths.append(f'温度範囲 {temp_text}')
    if base_product.get('description'):
        strengths.append(base_product['description'][:50])
    if not strengths:
        strengths.append(f'{base_maker}ブランド')

    # Weaknesses
    weaknesses = []
    if base_price and comp_prices:
        min_comp = min(comp_prices)
        if base_price > min_comp:
            weaknesses.append(f'最安競合比 +¥{base_price - min_comp:,.0f}')
    if not base_price:
        weaknesses.append('公開価格なし')
    if not weaknesses:
        weaknesses.append('改善ポイント要調査')

    # Opportunities
    opportunities = []
    if all_prices:
        sorted_p = sorted(all_prices)
        for i in range(1, len(sorted_p)):
            gap = sorted_p[i] - sorted_p[i - 1]
            if gap > 100000:
                opportunities.append(f'¥{sorted_p[i-1]:,.0f}〜¥{sorted_p[i]:,.0f}に空白帯')
                break
    opportunities.append('PBブランドでの差別化')

    # Threats
    threats = []
    threats.append(f'競合{len(category_products)}機種/{len(makers)}社')
    if comp_prices:
        threats.append(f'最安¥{min(comp_prices):,.0f}の低価格攻勢')

    return {
        'type': 'swot',
        'data': {
            'strengths': strengths[:4],
            'weaknesses': weaknesses[:4],
            'opportunities': opportunities[:4],
            'threats': threats[:4],
        }
    }


def _gen_positioning_visual(category_products, base_product, axis_x, axis_y):
    """ポジショニング散布図用データ"""
    positioning_data = _build_positioning_data(category_products, axis_x, axis_y)
    valid_points = [d for d in positioning_data
                    if d['x_value'] is not None and d['y_value'] is not None]

    by_maker = {}
    base_model = base_product.get('model', '')
    base_point = None

    for d in valid_points:
        maker = d['maker']
        if maker not in by_maker:
            by_maker[maker] = []
        by_maker[maker].append({
            'x': d['x_value'], 'y': d['y_value'], 'model': d['model']
        })
        if d['model'] == base_model:
            base_point = {'x': d['x_value'], 'y': d['y_value']}

    datasets = []
    for maker, points in by_maker.items():
        color = _MAKER_COLORS.get(maker, _DEFAULT_COLOR)
        datasets.append({
            'label': maker,
            'data': [{'x': p['x'], 'y': p['y']} for p in points],
            'models': [p['model'] for p in points],
            'backgroundColor': color['bg'],
            'borderColor': color['border'],
            'pointRadius': 6,
        })

    return {
        'type': 'positioning',
        'data': {
            'axis_x': axis_x,
            'axis_y': axis_y,
            'datasets': datasets,
            'base_product': base_point,
            'base_model': base_model,
        }
    }


def _gen_5forces_visual(category_products, base_product):
    """5Forces ダイヤモンド配置用データ"""
    makers = set(p.get('maker', '') for p in category_products)
    prices = [p.get('price_numeric') for p in category_products if p.get('price_numeric')]
    num_products = len(category_products)
    num_makers = len(makers)

    # 業界内競争
    if num_products >= 50:
        rivalry = 5
    elif num_products >= 30:
        rivalry = 4
    elif num_products >= 15:
        rivalry = 3
    elif num_products >= 5:
        rivalry = 2
    else:
        rivalry = 1

    # 新規参入の脅威（医療機器規制あり→低い）
    new_entrants = 2

    # 代替品の脅威（蒸気滅菌の代替は限定的）
    substitutes = 2

    # 買い手の交渉力
    if prices:
        price_range = max(prices) - min(prices)
        avg = sum(prices) / len(prices)
        cv = (price_range / avg) if avg > 0 else 0
        if cv > 1.0:
            buyer_power = 4
        elif cv > 0.5:
            buyer_power = 3
        else:
            buyer_power = 2
    else:
        buyer_power = 3

    # 売り手の交渉力
    if num_makers >= 5:
        supplier_power = 2
    elif num_makers >= 3:
        supplier_power = 3
    else:
        supplier_power = 4

    price_detail = ''
    if prices:
        price_detail = f'価格帯 ¥{min(prices):,.0f}〜¥{max(prices):,.0f}'

    return {
        'type': '5forces',
        'data': {
            'rivalry': {'score': rivalry, 'label': '業界内競争',
                        'detail': f'{num_products}機種/{num_makers}社'},
            'new_entrants': {'score': new_entrants, 'label': '新規参入の脅威',
                             'detail': '医療機器規制あり'},
            'substitutes': {'score': substitutes, 'label': '代替品の脅威',
                            'detail': '蒸気滅菌の代替は限定的'},
            'buyer_power': {'score': buyer_power, 'label': '買い手の交渉力',
                            'detail': price_detail},
            'supplier_power': {'score': supplier_power, 'label': '売り手の交渉力',
                               'detail': f'{num_makers}メーカー'},
        }
    }


def _gen_price_map_visual(category_products, base_product):
    """価格帯マップ横棒チャート用データ"""
    price_items = []
    for p in category_products:
        if p.get('price_numeric'):
            price_items.append({
                'maker': p.get('maker', ''),
                'model': p.get('model', ''),
                'price': p.get('price_numeric'),
            })
    price_items.sort(key=lambda x: x['price'])

    # 最大30件に制限（上位15+下位15）
    if len(price_items) > 30:
        price_items = price_items[:15] + price_items[-15:]

    labels = [f"{p['maker']}: {p['model']}" for p in price_items]
    data = [p['price'] for p in price_items]
    colors = [_MAKER_COLORS.get(p['maker'], _DEFAULT_COLOR)['bg'] for p in price_items]
    borders = [_MAKER_COLORS.get(p['maker'], _DEFAULT_COLOR)['border'] for p in price_items]

    return {
        'type': 'price_map',
        'data': {
            'labels': labels,
            'values': data,
            'colors': colors,
            'borders': borders,
            'base_price': base_product.get('price_numeric'),
            'base_model': base_product.get('model', ''),
        }
    }


def handle_analyze_framework(args, session):
    """フレームワーク分析（競合DBファクトベース）"""
    framework = args.get('framework')
    valid = ['3c', 'swot', 'positioning', '5forces', 'price_map']
    if framework not in valid:
        return {"error": f"無効なフレームワーク: {framework}", "valid": valid}

    axis_x = args.get('axis_x', '')
    axis_y = args.get('axis_y', '')

    # キャッシュキー（ポジショニングは軸込み）
    cache_key = framework
    if framework == 'positioning' and axis_x and axis_y:
        cache_key = f"positioning_{axis_x}_{axis_y}"
    # 分析コンテキスト構築
    base = session.get('base_product')
    if not base:
        maker_pn = session.get('pb_card', {}).get('maker_part_no')
        if maker_pn:
            db_results = _search_products(query=maker_pn)
            for p in db_results:
                if p.get('model') == maker_pn:
                    base = p
                    session['base_product'] = p
                    break
            if not base and db_results:
                base = db_results[0]
                session['base_product'] = base
    if not base:
        return {"error": "ベース製品が未選択です。先に製品を検索・選択してください。"}

    # 同カテゴリの全製品を取得（競合含む）
    category_products = _search_products(category=base.get('category'))
    pb_card = session.get('pb_card', {})
    competitor_summary = _build_competitor_summary(category_products, base)

    cached = session.get('framework_results', {}).get(cache_key)
    if cached:
        # キャッシュヒット時もビジュアルは生成
        visual = _generate_framework_visual(
            framework, category_products, base,
            axis_x=axis_x, axis_y=axis_y,
        )
        if visual:
            if '_pending_visuals' not in session:
                session['_pending_visuals'] = []
            session['_pending_visuals'].append(visual)
        return {"framework": framework, "result": cached, "cached": True}

    # --- ポジショニング: 軸未指定時は候補を返す ---
    if framework == 'positioning' and (not axis_x or not axis_y):
        candidates = _get_axis_candidates(category_products)
        return {
            "framework": "positioning",
            "needs_axes": True,
            "message": "ポジショニング分析の縦横軸を選んでください。",
            "axis_candidates": candidates,
            "suggested_pairs": [
                {"x": "価格", "y": "有効容量", "description": "価格 × 容量（コストパフォーマンス視点）"},
                {"x": "価格", "y": "使用温度範囲", "description": "価格 × 温度範囲（性能視点）"},
                {"x": "有効容量", "y": "重量", "description": "容量 × 重量（設置性視点）"},
            ],
            "competitor_count": len(category_products),
            "makers": list(set(p.get('maker', '') for p in category_products)),
        }

    # --- ポジショニング: 軸指定あり → 実データマップ生成 ---
    if framework == 'positioning':
        positioning_data = _build_positioning_data(category_products, axis_x, axis_y)
        valid_points = [d for d in positioning_data if d['x_value'] is not None and d['y_value'] is not None]

        prompt = f"""以下のポジショニングデータを元に「ポジショニングマップ分析」を実施してください。

## 軸設定
- X軸: {axis_x}
- Y軸: {axis_y}

## ベース製品（PB元）
{json.dumps(base, ensure_ascii=False, indent=2)}

## PBカード状態
{json.dumps(pb_card, ensure_ascii=False, indent=2)}

## ポジショニングデータ（{len(valid_points)}製品分・実データ）
{json.dumps(valid_points, ensure_ascii=False, indent=2)}

## データ不足の製品（{len(positioning_data) - len(valid_points)}件）
{json.dumps([d for d in positioning_data if d['x_value'] is None or d['y_value'] is None][:5], ensure_ascii=False, indent=2)}

以下の観点で分析してください:
1. テキストベースのポジショニングマップ（各象限に製品を配置）
2. 空白ゾーン（競合がいない価格帯・スペック帯）= PBのチャンス
3. ベース製品の現在位置と、PBとして狙うべきポジション
4. 具体的な差別化戦略の提案"""

    # --- 3C分析: 競合をファクトベースに ---
    elif framework == '3c':
        prompt = f"""以下の実データを元に「3C分析（Customer/Competitor/Company）」を実施してください。

## ベース製品（Company側の主力製品）
{json.dumps(base, ensure_ascii=False, indent=2)}

## PBカード状態
{json.dumps(pb_card, ensure_ascii=False, indent=2)}

## Competitor: 競合メーカーの実データ（DBに登録済み）
{competitor_summary}

## 同カテゴリ全製品（{len(category_products)}件中上位10件）
{json.dumps(category_products[:10], ensure_ascii=False, indent=2)}

【重要】Competitorセクションでは、上記の実データに基づいて分析してください。推測ではなくファクトベースで。
分析結果を構造化して、PB企画に活かせる具体的なインサイトを含めてください。"""

    # --- 価格帯マップ: 実データプロット ---
    elif framework == 'price_map':
        price_data = []
        for p in category_products:
            if p.get('price_numeric'):
                caps = p.get('specs', {}).get('有効容量', p.get('specs', {}).get('缶体有効内容積', ''))
                price_data.append({
                    'maker': p.get('maker'),
                    'model': p.get('model'),
                    'price': p.get('price_numeric'),
                    'capacity': caps,
                })
        price_data.sort(key=lambda x: x['price'])

        prompt = f"""以下の実データを元に「価格帯マップ分析」を実施してください。

## ベース製品
{json.dumps(base, ensure_ascii=False, indent=2)}

## PBカード状態
{json.dumps(pb_card, ensure_ascii=False, indent=2)}

## 価格データ（{len(price_data)}製品・実売価格）
{json.dumps(price_data, ensure_ascii=False, indent=2)}

## 競合サマリ
{competitor_summary}

以下の観点で分析してください:
1. メーカー別の価格帯分布（テキストベースのチャート）
2. 価格帯の空白ゾーン（競合がいない価格帯）= PBの参入余地
3. 容量あたり単価の比較（コスパランキング）
4. PB製品の推奨価格帯と根拠"""

    # --- SWOT/5Forces: 競合データ付き ---
    else:
        framework_names = {
            'swot': 'SWOT分析（Strengths/Weaknesses/Opportunities/Threats）',
            '5forces': '5Forces分析（業界構造分析）',
        }
        prompt = f"""以下の情報を元に「{framework_names[framework]}」を実施してください。

## ベース製品
{json.dumps(base, ensure_ascii=False, indent=2)}

## PBカード状態
{json.dumps(pb_card, ensure_ascii=False, indent=2)}

## 競合メーカーの実データ
{competitor_summary}

## 同カテゴリ全製品（{len(category_products)}件中上位10件）
{json.dumps(category_products[:10], ensure_ascii=False, indent=2)}

【重要】競合に関する記述は上記の実データに基づいてください。
分析結果を構造化して、PB企画に活かせる具体的なインサイトを含めてください。"""

    # Claude APIで分析実行
    analysis_result = _call_claude_simple(prompt)

    # キャッシュ保存
    if 'framework_results' not in session:
        session['framework_results'] = {}
    session['framework_results'][cache_key] = analysis_result

    # ビジュアルデータ生成
    visual = _generate_framework_visual(
        framework, category_products, base,
        axis_x=axis_x, axis_y=axis_y,
    )
    if visual:
        if '_pending_visuals' not in session:
            session['_pending_visuals'] = []
        session['_pending_visuals'].append(visual)

    return {"framework": framework, "result": analysis_result, "cached": False}


def handle_generate_pim_excel(args, session):
    """PIMデータExcel生成"""
    pb = session.get('pb_card', {})

    # spec_diffをspec_changesから自動生成（未設定の場合）
    spec_changes = session.get('spec_changes', [])
    if not pb.get('spec_diff') and spec_changes:
        pb['spec_diff'] = _generate_spec_diff_summary(spec_changes)

    # spec_diff以外の未確定項目をチェック
    unfilled = [k for k, v in pb.items() if v is None and k != 'spec_diff']
    if unfilled:
        return {"error": f"未確定項目があります: {', '.join(unfilled)}"}

    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

        wb = Workbook()
        ws = wb.active
        ws.title = "PIMデータ"

        # ヘッダースタイル
        header_fill = PatternFill(start_color="E60012", end_color="E60012", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        # ヘッダー
        headers = ["項目", "値"]
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border

        # データ
        field_labels = {
            'asone_part_no': 'アズワン品番',
            'price': '販売価格（税抜）',
            'jan_code': 'JANコード',
            'maker_part_no': 'メーカー型番',
            'quantity': '入数',
            'catchcopy': 'キャッチコピー',
            'spec_diff': '仕様差分',
        }

        for row_idx, (key, label) in enumerate(field_labels.items(), 2):
            ws.cell(row=row_idx, column=1, value=label).border = thin_border
            ws.cell(row=row_idx, column=2, value=pb.get(key, '')).border = thin_border

        # ベース製品情報
        base = session.get('base_product')
        if base:
            row = len(field_labels) + 3
            ws.cell(row=row, column=1, value="【ベース製品情報】").font = Font(bold=True)
            row += 1
            base_labels = {
                'name': '製品名', 'maker': 'メーカー', 'model': '型番',
                'category': 'カテゴリ', 'usage': '用途', 'price': '仕入れ先価格',
                'description': '概要',
            }
            for key, label in base_labels.items():
                if base.get(key):
                    ws.cell(row=row, column=1, value=label).border = thin_border
                    ws.cell(row=row, column=2, value=base[key]).border = thin_border
                    row += 1
            if base.get('specs'):
                for sk, sv in base['specs'].items():
                    ws.cell(row=row, column=1, value=sk).border = thin_border
                    ws.cell(row=row, column=2, value=str(sv)).border = thin_border
                    row += 1

        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 50

        filename = f"PIM_{pb.get('maker_part_no', 'unknown')}_{int(time.time())}.xlsx"
        filepath = os.path.join(_DOWNLOADS_DIR, filename)
        wb.save(filepath)
        return {"success": True, "filename": filename, "download_url": f"/api/download/{filename}"}

    except ImportError:
        return {"error": "openpyxlがインストールされていません"}
    except Exception as e:
        return {"error": f"Excel生成エラー: {str(e)}"}


def _generate_spec_diff_summary(spec_changes):
    """spec_changesから仕様差分サマリを自動生成"""
    if not spec_changes:
        return ""
    lines = []
    for ch in spec_changes:
        t = ch.get('type')
        if t == 'update':
            lines.append(f"変更: {ch.get('name')}「{ch.get('original')}」→「{ch.get('new')}」")
        elif t == 'delete':
            lines.append(f"削除: {ch.get('name')}「{ch.get('original')}」")
        elif t == 'add':
            lines.append(f"追加: {ch.get('name')}「{ch.get('new')}」")
    return "／".join(lines)


def handle_generate_proposal_word(args, session):
    """企画書Word生成（構成: 概要→フレームワーク→PB仕様→仕様諸元表）"""
    pb = session.get('pb_card', {})
    base = session.get('base_product')

    # ベース製品の自動復元
    if not base:
        maker_pn = pb.get('maker_part_no')
        if maker_pn:
            db_results = _search_products(query=maker_pn)
            for p in db_results:
                if p.get('model') == maker_pn:
                    base = p
                    session['base_product'] = p
                    break
            if not base and db_results:
                base = db_results[0]
                session['base_product'] = base

    fw = session.get('framework_results', {})
    spec_changes = session.get('spec_changes', [])

    # spec_diffをspec_changesから自動生成（PBカードに未設定の場合）
    if not pb.get('spec_diff') and spec_changes:
        pb['spec_diff'] = _generate_spec_diff_summary(spec_changes)

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        section_num = 1

        # タイトル
        title = doc.add_heading('PB企画書', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- 1. 企画概要 ---
        doc.add_heading(f'{section_num}. 企画概要', level=1)
        if base:
            doc.add_paragraph(f"ベース製品: {base.get('name', '')} ({base.get('maker', '')})")
            doc.add_paragraph(f"メーカー型番: {base.get('model', '')}")
            if base.get('price'):
                doc.add_paragraph(f"仕入れ先価格: {base.get('price')}")
            if base.get('usage'):
                doc.add_paragraph(f"用途: {base.get('usage')}")
            if base.get('description'):
                doc.add_paragraph(f"概要: {base.get('description')}")
        section_num += 1

        # --- 2. フレームワーク分析（前半） ---
        if fw:
            doc.add_heading(f'{section_num}. フレームワーク分析', level=1)
            fw_names = {
                '3c': '3C分析', 'swot': 'SWOT分析',
                'positioning': 'ポジショニング', '5forces': '5Forces',
                'price_map': '価格帯マップ'
            }
            for fw_key, fw_result in fw.items():
                doc.add_heading(fw_names.get(fw_key, fw_key), level=2)
                doc.add_paragraph(fw_result)
            section_num += 1

        # --- 3. PB製品仕様 ---
        doc.add_heading(f'{section_num}. PB製品仕様', level=1)
        field_labels = {
            'asone_part_no': 'アズワン品番',
            'price': '販売価格（税抜）',
            'jan_code': 'JANコード',
            'maker_part_no': 'メーカー型番',
            'quantity': '入数',
            'catchcopy': 'キャッチコピー',
        }
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '項目'
        hdr_cells[1].text = '値'
        for key, label in field_labels.items():
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(pb.get(key) or '—')
        # spec_diffがある場合は最後に追加
        if pb.get('spec_diff'):
            row = table.add_row().cells
            row[0].text = '仕様差分'
            row[1].text = str(pb['spec_diff'])
        section_num += 1

        # --- 4. 仕様諸元表（最終ページ） ---
        confirmed_specs = session.get('confirmed_specs', [])
        if confirmed_specs:
            # 改ページ
            p = doc.add_paragraph()
            run = p.add_run()
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run._element.append(br)

            doc.add_heading(f'{section_num}. 仕様諸元表', level=1)
            if base:
                doc.add_paragraph(
                    f"製品名: {base.get('name', '')}　型番: {base.get('model', '')}"
                )
            if spec_changes:
                doc.add_paragraph(
                    f"※ベースfixから{len(spec_changes)}件の変更を反映済み"
                )

            spec_tbl = doc.add_table(rows=1, cols=3)
            spec_tbl.style = 'Table Grid'
            hdr = spec_tbl.rows[0].cells
            hdr[0].text = 'No.'
            hdr[1].text = '項目'
            hdr[2].text = '仕様'
            for cell in hdr:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for item in confirmed_specs:
                row = spec_tbl.add_row().cells
                row[0].text = str(item.get('no', ''))
                row[1].text = str(item.get('name', ''))
                row[2].text = str(item.get('value', ''))

            for row in spec_tbl.rows:
                row.cells[0].width = Cm(1.5)
                row.cells[1].width = Cm(6)
                row.cells[2].width = Cm(10)

            # 変更履歴（spec_changesがある場合）
            if spec_changes:
                doc.add_heading('変更履歴', level=2)
                change_tbl = doc.add_table(rows=1, cols=4)
                change_tbl.style = 'Table Grid'
                ch_hdr = change_tbl.rows[0].cells
                ch_hdr[0].text = '種別'
                ch_hdr[1].text = '項目'
                ch_hdr[2].text = '内容'
                ch_hdr[3].text = '理由'
                for cell in ch_hdr:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                type_labels = {"update": "変更", "delete": "削除", "add": "追加"}
                for ch in spec_changes:
                    row = change_tbl.add_row().cells
                    row[0].text = type_labels.get(ch.get('type', ''), ch.get('type', ''))
                    row[1].text = str(ch.get('name', ''))
                    if ch['type'] == 'update':
                        row[2].text = f"{ch.get('original', '')} → {ch.get('new', '')}"
                    elif ch['type'] == 'delete':
                        row[2].text = f"{ch.get('original', '')}（削除）"
                    elif ch['type'] == 'add':
                        row[2].text = str(ch.get('new', ''))
                    row[3].text = str(ch.get('reason', ''))

        filename = f"企画書_{pb.get('maker_part_no', 'unknown')}_{int(time.time())}.docx"
        filepath = os.path.join(_DOWNLOADS_DIR, filename)
        doc.save(filepath)
        return {"success": True, "filename": filename, "download_url": f"/api/download/{filename}"}

    except ImportError:
        return {"error": "python-docxがインストールされていません"}
    except Exception as e:
        return {"error": f"Word生成エラー: {str(e)}"}


def handle_translate_to_english(args, session):
    """PIMデータ英訳Excel生成"""
    pb = session.get('pb_card', {})
    unfilled = [k for k, v in pb.items() if v is None]
    if unfilled:
        return {"error": f"未確定項目があります: {', '.join(unfilled)}"}

    base = session.get('base_product')

    # Claude APIで翻訳
    translate_prompt = f"""以下のPB製品情報を英訳してください。JSON形式で返してください。

PBカード:
{json.dumps(pb, ensure_ascii=False, indent=2)}

ベース製品:
{json.dumps(base, ensure_ascii=False, indent=2) if base else 'なし'}

出力形式（JSON）:
{{
  "asone_part_no": "...",
  "price": "...",
  "jan_code": "...",
  "maker_part_no": "...",
  "quantity": "...",
  "catchcopy": "...(英訳)",
  "spec_diff": "...(英訳)",
  "product_name": "...(英訳)",
  "description": "...(英訳)"
}}"""

    translated_text = _call_claude_simple(translate_prompt)

    # JSONパース試行
    try:
        # JSON部分を抽出
        json_match = re.search(r'\{[^{}]*\}', translated_text, re.DOTALL)
        if json_match:
            translated = json.loads(json_match.group())
        else:
            translated = {"raw_translation": translated_text}
    except json.JSONDecodeError:
        translated = {"raw_translation": translated_text}

    # Excel生成
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Border, Side, Alignment

        wb = Workbook()
        ws = wb.active
        ws.title = "PIM Data (English)"

        header_fill = PatternFill(start_color="E60012", end_color="E60012", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        headers = ["Field", "Japanese", "English"]
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border

        field_labels = {
            'asone_part_no': 'AS ONE Part No.',
            'price': 'Price (excl. tax)',
            'jan_code': 'JAN Code',
            'maker_part_no': 'Maker Part No.',
            'quantity': 'Quantity',
            'catchcopy': 'Catchcopy',
            'spec_diff': 'Spec Difference',
        }

        for row_idx, (key, label) in enumerate(field_labels.items(), 2):
            ws.cell(row=row_idx, column=1, value=label).border = thin_border
            ws.cell(row=row_idx, column=2, value=str(pb.get(key, ''))).border = thin_border
            ws.cell(row=row_idx, column=3, value=str(translated.get(key, ''))).border = thin_border

        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 40
        ws.column_dimensions['C'].width = 40

        filename = f"PIM_EN_{pb.get('maker_part_no', 'unknown')}_{int(time.time())}.xlsx"
        filepath = os.path.join(_DOWNLOADS_DIR, filename)
        wb.save(filepath)
        return {"success": True, "filename": filename, "download_url": f"/api/download/{filename}"}

    except Exception as e:
        return {"error": f"英訳Excel生成エラー: {str(e)}"}


def handle_generate_catalog_html(args, session):
    """カタログHTML生成"""
    pb = session.get('pb_card', {})
    base = session.get('base_product')

    # ベース製品の自動復元
    if not base:
        maker_pn = pb.get('maker_part_no')
        if maker_pn:
            db_results = _search_products(query=maker_pn)
            for p in db_results:
                if p.get('model') == maker_pn:
                    base = p
                    session['base_product'] = p
                    break
            if not base and db_results:
                base = db_results[0]
                session['base_product'] = base
    if not base:
        return {"error": "ベース製品が未選択です"}

    specs_html = ""
    if base.get('specs'):
        specs_rows = ""
        for k, v in base['specs'].items():
            specs_rows += f"<tr><td>{k}</td><td>{v}</td></tr>\n"
        specs_html = f"""
        <h2>仕様</h2>
        <table class="spec-table">
            <thead><tr><th>項目</th><th>値</th></tr></thead>
            <tbody>{specs_rows}</tbody>
        </table>"""

    diff_section = ""
    if pb.get('spec_diff'):
        diff_section = f"""
        <div class="diff-section">
            <h3>PB仕様差分</h3>
            <p>{pb['spec_diff']}</p>
        </div>"""

    html_content = f"""<!DOCTYPE html>
<html lang="ja">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{base.get('name', 'PB製品')} - カタログ</title>
<style>
body {{ font-family: 'Helvetica Neue', Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; color: #333; }}
h1 {{ color: #E60012; border-bottom: 3px solid #E60012; padding-bottom: 10px; }}
h2 {{ color: #E60012; margin-top: 30px; }}
.product-info {{ background: #F5F5F5; padding: 20px; border-radius: 8px; margin: 20px 0; }}
.product-info p {{ margin: 8px 0; }}
.spec-table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
.spec-table th, .spec-table td {{ border: 1px solid #E0E0E0; padding: 10px; text-align: left; }}
.spec-table th {{ background: #E60012; color: white; }}
.spec-table tr:nth-child(even) {{ background: #F9F9F9; }}
.catchcopy {{ font-size: 1.3em; color: #E60012; font-weight: bold; margin: 20px 0; padding: 15px; border-left: 4px solid #E60012; background: #FFF5F5; }}
.diff-section {{ background: #FFF8E1; padding: 15px; border-radius: 8px; margin: 20px 0; border-left: 4px solid #FFA000; }}
.pb-badge {{ display: inline-block; background: #E60012; color: white; padding: 4px 12px; border-radius: 4px; font-size: 12px; font-weight: bold; }}
.footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #E0E0E0; color: #666; font-size: 12px; }}
</style>
</head>
<body>
<h1><span class="pb-badge">PB</span> {base.get('name', '')}</h1>

{f'<div class="catchcopy">{pb["catchcopy"]}</div>' if pb.get('catchcopy') else ''}

<div class="product-info">
    <p><strong>アズワン品番:</strong> {pb.get('asone_part_no', '—')}</p>
    <p><strong>販売価格:</strong> {pb.get('price', '—')}</p>
    <p><strong>JANコード:</strong> {pb.get('jan_code', '—')}</p>
    <p><strong>メーカー型番:</strong> {pb.get('maker_part_no', '—')}</p>
    <p><strong>入数:</strong> {pb.get('quantity', '—')}</p>
    <p><strong>メーカー:</strong> {base.get('maker', '—')}</p>
    <p><strong>用途:</strong> {base.get('usage', '—')}</p>
</div>

{diff_section}
{specs_html}

<div class="footer">
    <p>Generated by PB Planner | AS ONE Corporation</p>
</div>
</body>
</html>"""

    filename = f"catalog_{pb.get('maker_part_no', 'unknown')}_{int(time.time())}.html"
    filepath = os.path.join(_DOWNLOADS_DIR, filename)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(html_content)

    return {"success": True, "filename": filename, "download_url": f"/api/download/{filename}"}


def handle_save_spec_list(args, session):
    """仕様諸元リストを保存"""
    specs = args.get('specs', [])
    if not specs:
        return {"error": "仕様リストが空です"}

    # 連番を正規化して保存
    confirmed = []
    for i, item in enumerate(specs):
        confirmed.append({
            "no": item.get('no', i + 1),
            "name": str(item.get('name', '')),
            "value": str(item.get('value', '')),
        })
    session['confirmed_specs'] = confirmed
    return {
        "success": True,
        "count": len(confirmed),
        "message": f"仕様諸元リスト {len(confirmed)}項目を保存しました。企画書Wordの最終ページに仕様諸元表として挿入されます。"
    }


def handle_update_spec_item(args, session):
    """仕様諸元リストの特定番号の項目を更新 + 変更ログ記録"""
    no = args.get('no')
    new_value = args.get('value')
    new_name = args.get('name')
    reason = args.get('reason', '')

    specs = session.get('confirmed_specs', [])
    if not specs:
        return {"error": "仕様諸元リストが未保存です。先にsave_spec_listで保存してください。"}

    # 番号で検索
    found_item = None
    for item in specs:
        if item['no'] == no:
            found_item = item
            break

    if not found_item:
        return {"error": f"番号 {no} の項目が見つかりません（範囲: 1〜{len(specs)}）"}

    # 変更ログ記録
    change_log = {
        "type": "update",
        "no": no,
        "name": found_item['name'],
        "original": found_item['value'],
        "new": new_value,
        "reason": reason,
    }
    if new_name:
        change_log["original_name"] = found_item['name']
        change_log["new_name"] = new_name

    if 'spec_changes' not in session:
        session['spec_changes'] = []
    session['spec_changes'].append(change_log)

    # 実際の更新
    found_item['value'] = new_value
    if new_name:
        found_item['name'] = new_name

    session['confirmed_specs'] = specs
    return {
        "success": True,
        "updated": {"no": no, "name": new_name or found_item['name'], "value": new_value},
        "change_count": len(session['spec_changes']),
        "message": f"{no}番の値を「{new_value}」に更新しました。（変更ログ記録済み）"
    }


def handle_delete_spec_item(args, session):
    """仕様諸元リストから特定番号の項目を削除 + 変更ログ記録"""
    no = args.get('no')
    reason = args.get('reason', '')

    specs = session.get('confirmed_specs', [])
    if not specs:
        return {"error": "仕様諸元リストが未保存です。先にsave_spec_listで保存してください。"}

    # 番号で検索
    found_item = None
    found_idx = None
    for idx, item in enumerate(specs):
        if item['no'] == no:
            found_item = item
            found_idx = idx
            break

    if found_item is None:
        return {"error": f"番号 {no} の項目が見つかりません（範囲: 1〜{len(specs)}）"}

    # 変更ログ記録
    if 'spec_changes' not in session:
        session['spec_changes'] = []
    session['spec_changes'].append({
        "type": "delete",
        "no": no,
        "name": found_item['name'],
        "original": found_item['value'],
        "reason": reason,
    })

    # 削除実行
    specs.pop(found_idx)
    session['confirmed_specs'] = specs
    return {
        "success": True,
        "deleted": {"no": no, "name": found_item['name']},
        "remaining_count": len(specs),
        "change_count": len(session['spec_changes']),
        "message": f"{no}番「{found_item['name']}」を削除しました。（変更ログ記録済み）"
    }


def handle_add_spec_item(args, session):
    """仕様諸元リストに新しい項目を追加 + 変更ログ記録"""
    name = args.get('name', '')
    value = args.get('value', '')
    reason = args.get('reason', '')

    specs = session.get('confirmed_specs', [])
    if not specs:
        return {"error": "仕様諸元リストが未保存です。先にsave_spec_listで保存してください。"}

    # 新しい番号（最大値+1）
    max_no = max((item['no'] for item in specs), default=0)
    new_no = max_no + 1

    new_item = {"no": new_no, "name": name, "value": value}
    specs.append(new_item)

    # 変更ログ記録
    if 'spec_changes' not in session:
        session['spec_changes'] = []
    session['spec_changes'].append({
        "type": "add",
        "no": new_no,
        "name": name,
        "new": value,
        "reason": reason,
    })

    session['confirmed_specs'] = specs
    return {
        "success": True,
        "added": {"no": new_no, "name": name, "value": value},
        "total_count": len(specs),
        "change_count": len(session['spec_changes']),
        "message": f"{new_no}番「{name}: {value}」を追加しました。（変更ログ記録済み）"
    }


# ================================================================
# Claude API 呼び出し
# ================================================================

def _call_claude_simple(prompt, max_tokens=2000):
    """シンプルなClaude API呼び出し（ツールなし）"""
    import urllib.request

    payload = {
        "model": "claude-sonnet-4-20250514",
        "max_tokens": max_tokens,
        "messages": [{"role": "user", "content": prompt}],
    }

    data = json.dumps(payload).encode('utf-8')
    req = urllib.request.Request(
        'https://api.anthropic.com/v1/messages',
        data=data,
        headers={
            'Content-Type': 'application/json',
            'x-api-key': CLAUDE_API_KEY,
            'anthropic-version': '2023-06-01',
        }
    )

    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            result = json.loads(resp.read().decode('utf-8'))
            for block in result.get('content', []):
                if block.get('type') == 'text':
                    return block['text']
            return ''
    except Exception as e:
        return f"[API Error: {str(e)}]"


def _call_claude_with_tools(messages, system_prompt, session):
    """Claude API呼び出し（Function Calling付き）。download_urlを含む結果はリストで返す。"""
    import urllib.request

    payload = {
        "model": "claude-sonnet-4-20250514",
        "max_tokens": 4096,
        "system": system_prompt,
        "messages": messages,
        "tools": FC_TOOLS,
    }

    all_text_parts = []
    download_urls = []

    max_iterations = 10
    for iteration in range(max_iterations):
        data = json.dumps(payload).encode('utf-8')
        req = urllib.request.Request(
            'https://api.anthropic.com/v1/messages',
            data=data,
            headers={
                'Content-Type': 'application/json',
                'x-api-key': CLAUDE_API_KEY,
                'anthropic-version': '2023-06-01',
            }
        )

        try:
            with urllib.request.urlopen(req, timeout=120) as resp:
                result = json.loads(resp.read().decode('utf-8'))
        except Exception as e:
            return f"API呼び出しエラー: {str(e)}", download_urls

        # レスポンス解析
        text_parts = []
        tool_calls = []

        for block in result.get('content', []):
            if block.get('type') == 'text':
                text_parts.append(block['text'])
            elif block.get('type') == 'tool_use':
                tool_calls.append(block)

        # ツール呼び出しがなければ完了
        if not tool_calls:
            all_text_parts.extend(text_parts)
            return '\n'.join(all_text_parts), download_urls

        # 中間テキストも保持
        all_text_parts.extend(text_parts)

        # ツール実行
        tool_use_blocks = result.get('content', [])
        messages.append({"role": "assistant", "content": tool_use_blocks})

        tool_results = []
        for tc in tool_calls:
            tool_name = tc['name']
            tool_input = tc.get('input', {})
            tool_id = tc['id']

            print(f"[TOOL] Calling: {tool_name}({json.dumps(tool_input, ensure_ascii=False)[:200]})", flush=True)
            handler = _TOOL_HANDLERS.get(tool_name)
            if handler:
                try:
                    result_data = handler(tool_input, session)
                    pending = session.get('_pending_visuals', [])
                    print(f"[TOOL] {tool_name} done. _pending_visuals count: {len(pending)}", flush=True)
                except Exception as e:
                    result_data = {"error": str(e)}
                    print(f"[TOOL] {tool_name} ERROR: {e}", flush=True)
            else:
                result_data = {"error": f"Unknown tool: {tool_name}"}

            # download_url を収集
            if isinstance(result_data, dict) and result_data.get('download_url'):
                download_urls.append({
                    "tool": tool_name,
                    "filename": result_data.get('filename', ''),
                    "download_url": result_data['download_url'],
                })

            tool_results.append({
                "type": "tool_result",
                "tool_use_id": tool_id,
                "content": json.dumps(result_data, ensure_ascii=False),
            })

        messages.append({"role": "user", "content": tool_results})
        payload["messages"] = messages

    return "最大反復回数に達しました。" + '\n'.join(all_text_parts), download_urls


# ツールハンドラマッピング
_TOOL_HANDLERS = {
    'search_products': handle_search_products,
    'set_pb_field': handle_set_pb_field,
    'get_pb_card': handle_get_pb_card,
    'analyze_framework': handle_analyze_framework,
    'generate_pim_excel': handle_generate_pim_excel,
    'generate_proposal_word': handle_generate_proposal_word,
    'translate_to_english': handle_translate_to_english,
    'generate_catalog_html': handle_generate_catalog_html,
    'save_spec_list': handle_save_spec_list,
    'update_spec_item': handle_update_spec_item,
    'delete_spec_item': handle_delete_spec_item,
    'add_spec_item': handle_add_spec_item,
}


# ================================================================
# ルート定義
# ================================================================

@app.route('/')
def index():
    return send_from_directory('static', 'index.html')


@app.route('/api/health')
def health():
    return jsonify({
        "status": "ok",
        "service": "pb-planner",
        "api_key_set": bool(CLAUDE_API_KEY),
        "api_key_prefix": CLAUDE_API_KEY[:10] + '...' if CLAUDE_API_KEY else 'NOT SET',
    })


@app.route('/api/chat', methods=['POST'])
def chat():
    """PB企画チャットエンドポイント"""
    body = request.get_json(silent=True) or {}
    user_msg = (body.get('message') or '').strip()
    session_id = body.get('session_id') or f'web_{int(time.time())}_{os.urandom(4).hex()}'

    if not user_msg:
        return jsonify({"error": "メッセージが空です"}), 400

    if not CLAUDE_API_KEY:
        return jsonify({"error": "APIキーが設定されていません"}), 500

    session = get_or_create_session(session_id)

    # base_product自動復元（ワーカー切替・セッション消失対策）
    if not session.get('base_product'):
        maker_pn = session.get('pb_card', {}).get('maker_part_no')
        if maker_pn:
            db_results = _search_products(query=maker_pn)
            for p in db_results:
                if p.get('model') == maker_pn:
                    session['base_product'] = p
                    break
            if not session.get('base_product') and db_results:
                session['base_product'] = db_results[0]

    # 会話履歴にユーザーメッセージ追加
    session['history'].append({"role": "user", "content": user_msg})

    # 履歴を最新20ターンに制限
    if len(session['history']) > 40:
        session['history'] = session['history'][-40:]

    # システムプロンプト構築
    system_prompt = _build_system_prompt(session)

    # Claude API呼び出し
    reply_text, download_urls = _call_claude_with_tools(
        list(session['history']),  # コピーを渡す
        system_prompt,
        session,
    )

    # AIの返答を履歴に追加
    session['history'].append({"role": "assistant", "content": reply_text})

    # spec_diffをspec_changesから自動生成（save_sessionの前に実行）
    spec_changes = session.get('spec_changes', [])
    if spec_changes:
        pb_card = session.get('pb_card', {})
        pb_card['spec_diff'] = _generate_spec_diff_summary(spec_changes)
        session['pb_card'] = pb_card

    # セッション保存（spec_diff反映後に保存）
    save_session(session_id, session)

    # レスポンス構築
    response = {
        "reply": reply_text,
        "session_id": session_id,
        "pb_card": session.get('pb_card', {}),
        "base_product": session.get('base_product'),
        "framework_results": list(session.get('framework_results', {}).keys()),
        "framework_visuals": session.pop('_pending_visuals', []),
        "download_urls": download_urls,
        "confirmed_specs_count": len(session.get('confirmed_specs', [])),
        "spec_changes_count": len(session.get('spec_changes', [])),
    }

    return jsonify(response)


@app.route('/api/reset', methods=['POST'])
def api_reset():
    """セッションリセット"""
    body = request.get_json(silent=True) or {}
    session_id = body.get('session_id', '')
    if session_id:
        if _redis_client:
            try:
                _redis_client.delete(f'pb:{session_id}')
            except Exception:
                pass
        elif session_id in _SESSIONS:
            del _SESSIONS[session_id]
    return jsonify({"success": True})


@app.route('/api/download/<filename>')
def download_file(filename):
    """生成ファイルダウンロード"""
    # パストラバーサル防止
    safe_name = os.path.basename(filename)
    filepath = os.path.join(_DOWNLOADS_DIR, safe_name)
    if not os.path.exists(filepath):
        abort(404)
    return send_file(filepath, as_attachment=True, download_name=safe_name)


@app.route('/<path:filename>')
def static_files(filename):
    """静的ファイル配信"""
    return send_from_directory('static', filename)


# ================================================================
# エントリポイント
# ================================================================

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
