"""
PB企画書Word生成（フレームワーク分析フル版）
- 3C分析 / SWOT / ポジショニング / 5Forces / 価格帯マップ
- 競合データ（81製品）に基づくデータ駆動分析
"""
import json
import os
import sys
import time

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ================================================================
# データ収集
# ================================================================
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, 'workspace', 'data')

def load_all_autoclaves():
    products = []
    for root, dirs, files in os.walk(DATA_DIR):
        for f in files:
            if f == 'products.jsonl':
                path = os.path.join(root, f)
                with open(path, 'r', encoding='utf-8') as fh:
                    for line in fh:
                        try:
                            p = json.loads(line.strip())
                            if ('autoclave' in root or
                                p.get('category', '') == 'autoclave' or
                                'オートクレーブ' in p.get('name', '') or
                                '滅菌' in p.get('name', '')):
                                products.append(p)
                        except:
                            pass
    return products

def find_base_product(products, model='FLS-1000'):
    for p in products:
        if p.get('model') == model:
            return p
    return None

def maker_summary(products):
    makers = {}
    for p in products:
        m = p.get('maker', '不明')
        if m not in makers:
            makers[m] = {'count': 0, 'prices': [], 'models': []}
        makers[m]['count'] += 1
        makers[m]['models'].append(p.get('model', ''))
        pn = p.get('price_numeric')
        if pn:
            makers[m]['prices'].append(pn)
    return makers

def extract_numeric(specs, keys):
    """specsから数値を抽出"""
    if not specs:
        return None
    for k in keys:
        for sk, sv in specs.items():
            if k in sk:
                import re
                nums = re.findall(r'[\d,.]+', str(sv))
                if nums:
                    try:
                        return float(nums[0].replace(',', ''))
                    except:
                        pass
    return None

# ================================================================
# フレームワーク分析テキスト生成
# ================================================================

def generate_3c(base, makers, all_products):
    lines = []
    lines.append("【Company（自社）】")
    lines.append(f"・ベース製品: {base.get('name', '')} ({base.get('model', '')})")
    lines.append(f"・メーカー: {base.get('maker', '')}")
    if base.get('price'):
        lines.append(f"・仕入れ先価格: {base['price']}")
    if base.get('design_concept'):
        lines.append(f"・設計コンセプト: {base['design_concept'][:100]}")
    features = base.get('features', {})
    if features:
        lines.append("・主要特長:")
        if isinstance(features, dict):
            for k, v in list(features.items())[:5]:
                lines.append(f"  - {k}: {str(v)[:60]}")
        elif isinstance(features, list):
            for feat in features[:5]:
                if isinstance(feat, dict):
                    lines.append(f"  - {feat.get('name', '')}: {feat.get('detail', '')[:60]}")
                else:
                    lines.append(f"  - {str(feat)[:80]}")
    specs = base.get('specs', {})
    if specs:
        lines.append("・主要スペック:")
        for k, v in list(specs.items())[:8]:
            lines.append(f"  - {k}: {v}")

    lines.append("")
    lines.append("【Competitor（競合）】")
    for m, data in makers.items():
        if m == base.get('maker'):
            continue
        prices = data['prices']
        price_str = f"（{min(prices):,.0f}〜{max(prices):,.0f}円）" if prices else "（価格非公開）"
        lines.append(f"・{m}: {data['count']}機種 {price_str}")
        lines.append(f"  代表型番: {', '.join(data['models'][:5])}")

    lines.append("")
    lines.append("【Customer（顧客）】")
    lines.append("・主要ターゲット: 研究機関、医療施設、食品検査機関")
    lines.append("・ニーズ: 大容量・省スペース・操作性・安全性・滅菌信頼性")
    usage = base.get('usage', '')
    if usage:
        lines.append(f"・ベース製品の用途: {usage}")

    return "\n".join(lines)

def generate_swot(base, makers, all_products):
    # 競合価格集計
    comp_prices = []
    for p in all_products:
        if p.get('maker') != base.get('maker') and p.get('price_numeric'):
            comp_prices.append(p['price_numeric'])

    base_price = 850000  # PB価格

    lines = []
    lines.append("【Strengths（強み）】")
    features = base.get('features', {})
    if isinstance(features, dict):
        for k, v in list(features.items())[:4]:
            lines.append(f"・{k}: {str(v)[:50]}")
    elif isinstance(features, list):
        for feat in features[:4]:
            if isinstance(feat, dict):
                lines.append(f"・{feat.get('name', '')}")
            else:
                lines.append(f"・{str(feat)[:60]}")
    specs = base.get('specs', {})
    for k, v in specs.items():
        if '特許' in str(v) or 'φ45' in str(v):
            lines.append(f"・{k}: {v}")

    lines.append("")
    lines.append("【Weaknesses（弱み）】")
    if comp_prices:
        min_comp = min(comp_prices)
        lines.append(f"・最安競合（{min_comp:,.0f}円）との価格差: +{base_price - min_comp:,.0f}円")
    lines.append("・ブランド認知度（PB新規参入のため）")
    lines.append("・アフターサービス体制の構築が必要")

    lines.append("")
    lines.append("【Opportunities（機会）】")
    if comp_prices:
        # 価格帯空白分析
        sorted_prices = sorted(comp_prices)
        gaps = []
        for i in range(len(sorted_prices) - 1):
            gap = sorted_prices[i+1] - sorted_prices[i]
            if gap > 100000:
                gaps.append((sorted_prices[i], sorted_prices[i+1]))
        if gaps:
            g = gaps[0]
            lines.append(f"・価格帯空白ゾーン: {g[0]:,.0f}〜{g[1]:,.0f}円")
    lines.append("・ラボ・医療分野の設備更新需要")
    lines.append("・女性研究者増加に伴う操作性重視トレンド")
    lines.append("・大容量（100L超）市場の競合が少ない")

    lines.append("")
    lines.append("【Threats（脅威）】")
    total_comp = sum(d['count'] for m, d in makers.items() if m != base.get('maker'))
    lines.append(f"・競合製品数: {total_comp}機種（{len(makers)-1}社）")
    lines.append("・海外メーカーの低価格攻勢")
    lines.append("・技術的参入障壁の低下")

    return "\n".join(lines)

def generate_positioning(base, all_products):
    lines = []
    lines.append("【ポジショニングマップ分析】")
    lines.append("軸: X=価格（円）、Y=缶体容量（L）")
    lines.append("")

    # データ収集
    data_points = []
    for p in all_products:
        price = p.get('price_numeric')
        capacity = extract_numeric(p.get('specs', {}), ['缶体容量', '容量', '有効容量'])
        if price and capacity:
            data_points.append({
                'maker': p.get('maker', ''),
                'model': p.get('model', ''),
                'price': price,
                'capacity': capacity
            })

    # メーカー別にプロット情報
    by_maker = {}
    for dp in data_points:
        m = dp['maker']
        if m not in by_maker:
            by_maker[m] = []
        by_maker[m].append(dp)

    for m, pts in by_maker.items():
        lines.append(f"■ {m}:")
        for pt in sorted(pts, key=lambda x: x['price']):
            lines.append(f"  {pt['model']}: {pt['price']:,.0f}円 / {pt['capacity']:.0f}L")

    lines.append("")
    lines.append(f"■ PB製品（★）: AS-FLS100: 850,000円 / 120L")

    lines.append("")
    lines.append("【インサイト】")
    # 近い競合を探す
    if data_points:
        lines.append(f"・分析対象: {len(data_points)}製品")
        nearby = [dp for dp in data_points if abs(dp['price'] - 850000) < 300000]
        if nearby:
            lines.append(f"・価格帯近接競合（±30万円）: {len(nearby)}製品")
            for dp in nearby[:3]:
                lines.append(f"  - {dp['maker']} {dp['model']}: {dp['price']:,.0f}円/{dp['capacity']:.0f}L")

        # 120L以上の競合
        large = [dp for dp in data_points if dp['capacity'] >= 100]
        lines.append(f"・大容量（100L以上）: {len(large)}製品")
        if not large:
            lines.append("  → 100L超は空白市場。PB製品120Lの差別化ポイント")

    return "\n".join(lines)

def generate_5forces(base, makers, all_products):
    total_makers = len(makers)
    total_products = len(all_products)
    comp_makers = total_makers - 1
    comp_products = total_products - len([p for p in all_products if p.get('maker') == base.get('maker')])

    # 価格分散
    prices = [p.get('price_numeric', 0) for p in all_products if p.get('price_numeric')]
    price_std = 0
    if prices:
        avg = sum(prices) / len(prices)
        price_std = (sum((p - avg)**2 for p in prices) / len(prices)) ** 0.5

    lines = []
    lines.append("【5Forces分析】")
    lines.append("")

    # 業界内競争
    rivalry = min(5, max(1, comp_products // 15))
    lines.append(f"■ 業界内競争の激しさ: {'●' * rivalry}{'○' * (5 - rivalry)} ({rivalry}/5)")
    lines.append(f"  競合{comp_makers}社・{comp_products}機種が参入")
    lines.append(f"  価格帯: {min(prices):,.0f}〜{max(prices):,.0f}円（標準偏差{price_std:,.0f}円）" if prices else "  価格データ不足")
    lines.append("")

    # 新規参入の脅威
    entry = 2
    lines.append(f"■ 新規参入の脅威: {'●' * entry}{'○' * (5 - entry)} ({entry}/5)")
    lines.append("  高圧容器は安全基準（小型圧力容器・労安法）の認証が必要")
    lines.append("  技術的参入障壁は中程度だが、品質実績の構築に時間を要する")
    lines.append("")

    # 代替品の脅威
    substitute = 2
    lines.append(f"■ 代替品の脅威: {'●' * substitute}{'○' * (5 - substitute)} ({substitute}/5)")
    lines.append("  EO滅菌、過酸化水素プラズマ滅菌など代替技術あり")
    lines.append("  ただし高圧蒸気滅菌は最も汎用的でコスト効率が高い")
    lines.append("")

    # 買い手の交渉力
    buyer = 4
    lines.append(f"■ 買い手の交渉力: {'●' * buyer}{'○' * (5 - buyer)} ({buyer}/5)")
    lines.append("  研究機関・病院は価格に敏感（予算制約）")
    lines.append("  アズワンカタログ掲載による比較容易性")
    lines.append("")

    # 売り手の交渉力
    supplier = 3
    lines.append(f"■ 売り手の交渉力: {'●' * supplier}{'○' * (5 - supplier)} ({supplier}/5)")
    lines.append(f"  仕入れ先候補: {total_makers}社")
    lines.append("  特殊部品（圧力容器本体）の調達先は限定的")

    return "\n".join(lines)

def generate_price_map(base, all_products):
    lines = []
    lines.append("【価格帯マップ】")
    lines.append("")

    # メーカー別・価格帯別集計
    bands = {
        '〜30万円': (0, 300000),
        '30〜50万円': (300000, 500000),
        '50〜80万円': (500000, 800000),
        '80〜100万円': (800000, 1000000),
        '100〜150万円': (1000000, 1500000),
        '150万円〜': (1500000, float('inf')),
    }

    by_maker = {}
    for p in all_products:
        m = p.get('maker', '不明')
        pn = p.get('price_numeric')
        if pn:
            if m not in by_maker:
                by_maker[m] = []
            by_maker[m].append({'model': p.get('model', ''), 'price': pn})

    # 価格帯別テーブル
    lines.append(f"{'価格帯':<12} ", )
    header = f"{'価格帯':<14}"
    for m in by_maker:
        header += f" {m:<10}"
    lines[-1] = header

    for band_name, (lo, hi) in bands.items():
        row = f"{band_name:<14}"
        for m in by_maker:
            count = sum(1 for p in by_maker[m] if lo <= p['price'] < hi)
            row += f" {count if count else '-':>10}"
        lines.append(row)

    lines.append("")
    lines.append("【製品別価格一覧】")
    all_priced = []
    for p in all_products:
        if p.get('price_numeric'):
            all_priced.append(p)
    all_priced.sort(key=lambda x: x['price_numeric'])

    for p in all_priced:
        marker = " ★" if p.get('model') == base.get('model') else ""
        lines.append(f"  {p['maker']:<8} {p.get('model', ''):<16} {p['price_numeric']:>10,.0f}円{marker}")

    lines.append("")
    lines.append(f"  ★ PB製品 AS-FLS100                850,000円 ← ターゲット価格")

    lines.append("")
    lines.append("【価格戦略サマリ】")
    if all_priced:
        avg = sum(p['price_numeric'] for p in all_priced) / len(all_priced)
        lines.append(f"・全製品平均価格: {avg:,.0f}円")
        lines.append(f"・PB価格 850,000円は平均{'以下' if 850000 <= avg else '以上'}")
        below = sum(1 for p in all_priced if p['price_numeric'] < 850000)
        above = sum(1 for p in all_priced if p['price_numeric'] >= 850000)
        lines.append(f"・PBより安い製品: {below}機種 / 高い製品: {above}機種")

    return "\n".join(lines)

# ================================================================
# Word生成
# ================================================================

def add_colored_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def add_framework_table(doc, title, content_text):
    """フレームワーク分析セクションを追加"""
    doc.add_heading(title, level=2)
    for line in content_text.split('\n'):
        if line.startswith('【') and line.endswith('】'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xE6, 0x00, 0x12)
        elif line.startswith('■'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(10)
        elif line.strip().startswith('・') or line.strip().startswith('-') or line.strip().startswith('→'):
            doc.add_paragraph(line.strip(), style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line.strip())

def generate_full_proposal():
    # データ読み込み
    all_products = load_all_autoclaves()
    base = find_base_product(all_products, 'FLS-1000')
    makers = maker_summary(all_products)

    if not base:
        print("ERROR: FLS-1000が見つかりません")
        return

    print(f"ベース製品: {base['name']} ({base['model']})")
    print(f"全製品数: {len(all_products)} ({len(makers)}社)")

    # PBカード
    pb_card = {
        'asone_part_no': 'AS-FLS100',
        'price': '850,000円',
        'jan_code': '4560000000001',
        'maker_part_no': 'FLS-1000',
        'quantity': '1台',
        'catchcopy': '大容量120L、片手で楽々操作',
        'spec_diff': '変更: 缶体容量「100リットル」→「120リットル」',
    }

    spec_changes = [{'type': 'update', 'name': '缶体容量',
                     'original': '100リットル', 'new': '120リットル', 'reason': 'PB差別化（大容量訴求）'}]

    # 仕様リスト
    confirmed_specs = []
    if base.get('specs'):
        for i, (k, v) in enumerate(base['specs'].items(), 1):
            val = str(v)
            if '缶体容量' in k or (k == '容量' and 'リットル' in val):
                val = '120リットル'
            confirmed_specs.append({'no': i, 'name': k, 'value': val})

    # フレームワーク分析生成
    fw_3c = generate_3c(base, makers, all_products)
    fw_swot = generate_swot(base, makers, all_products)
    fw_positioning = generate_positioning(base, all_products)
    fw_5forces = generate_5forces(base, makers, all_products)
    fw_price_map = generate_price_map(base, all_products)

    # ==========================================
    # Word Document
    # ==========================================
    doc = Document()

    # デフォルトフォント
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic'
    font.size = Pt(10)

    section_num = 1

    # ===== 表紙 =====
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_heading('PB企画書', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0xE6, 0x00, 0x12)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'{base["name"]}（{base["model"]}）ベース')
    run.font.size = Pt(14)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('2026年3月18日')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 改ページ
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

    # ===== 1. 企画概要 =====
    doc.add_heading(f'{section_num}. 企画概要', level=1)
    doc.add_paragraph(f"ベース製品: {base.get('name', '')} ({base.get('maker', '')})")
    doc.add_paragraph(f"メーカー型番: {base.get('model', '')}")
    if base.get('price'):
        doc.add_paragraph(f"仕入れ先価格: {base['price']}")
    if base.get('usage'):
        doc.add_paragraph(f"用途: {base['usage']}")
    if base.get('description'):
        doc.add_paragraph(f"概要: {base['description']}")
    if base.get('design_concept'):
        doc.add_paragraph(f"設計コンセプト: {base['design_concept']}")
    section_num += 1

    # ===== 2. PB製品仕様 =====
    doc.add_heading(f'{section_num}. PB製品仕様', level=1)
    field_labels = {
        'asone_part_no': 'アズワン品番',
        'price': '販売価格（税抜）',
        'jan_code': 'JANコード',
        'maker_part_no': 'メーカー型番（ベース）',
        'quantity': '入数',
        'catchcopy': 'キャッチコピー',
        'spec_diff': '仕様差分',
    }
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '項目'
    hdr_cells[1].text = '値'
    for p_run in hdr_cells[0].paragraphs[0].runs:
        p_run.bold = True
    for p_run in hdr_cells[1].paragraphs[0].runs:
        p_run.bold = True

    for key, label in field_labels.items():
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(pb_card.get(key, '—'))
    section_num += 1

    # 改ページ
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

    # ===== 3. フレームワーク分析 =====
    doc.add_heading(f'{section_num}. フレームワーク分析', level=1)
    doc.add_paragraph(f"競合データ: {len(all_products)}製品（{len(makers)}社）に基づく分析")

    add_framework_table(doc, '3-1. 3C分析', fw_3c)

    # 改ページ
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

    add_framework_table(doc, '3-2. SWOT分析', fw_swot)

    # 改ページ
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

    add_framework_table(doc, '3-3. ポジショニングマップ', fw_positioning)

    # 改ページ
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

    add_framework_table(doc, '3-4. 5Forces分析', fw_5forces)

    # 改ページ
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

    add_framework_table(doc, '3-5. 価格帯マップ', fw_price_map)
    section_num += 1

    # 改ページ
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

    # ===== 4. 仕様諸元表 =====
    if confirmed_specs:
        doc.add_heading(f'{section_num}. 仕様諸元表', level=1)
        doc.add_paragraph(f"製品名: {base.get('name', '')}　型番: {base.get('model', '')}")
        doc.add_paragraph(f"※ベースfixから{len(spec_changes)}件の変更を反映済み")

        spec_tbl = doc.add_table(rows=1, cols=3)
        spec_tbl.style = 'Table Grid'
        h = spec_tbl.rows[0].cells
        h[0].text = 'No.'
        h[1].text = '項目'
        h[2].text = '仕様'
        for cell in h:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for item in confirmed_specs:
            row = spec_tbl.add_row().cells
            row[0].text = str(item['no'])
            row[1].text = item['name']
            row[2].text = item['value']

        for row in spec_tbl.rows:
            row.cells[0].width = Cm(1.5)
            row.cells[1].width = Cm(6)
            row.cells[2].width = Cm(10)

        # 変更履歴
        doc.add_heading('変更履歴', level=2)
        ch_tbl = doc.add_table(rows=1, cols=4)
        ch_tbl.style = 'Table Grid'
        ch_hdr = ch_tbl.rows[0].cells
        ch_hdr[0].text = '種別'
        ch_hdr[1].text = '項目'
        ch_hdr[2].text = '内容'
        ch_hdr[3].text = '理由'
        for cell in ch_hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for ch in spec_changes:
            row = ch_tbl.add_row().cells
            row[0].text = '変更'
            row[1].text = ch['name']
            row[2].text = f"{ch['original']} → {ch['new']}"
            row[3].text = ch['reason']

    # 保存
    output_dir = os.path.join(BASE_DIR, 'downloads')
    os.makedirs(output_dir, exist_ok=True)
    filename = f"企画書_FLS-1000_フル版_{int(time.time())}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"\n✅ 企画書生成完了: {filepath}")
    return filepath

if __name__ == '__main__':
    generate_full_proposal()
